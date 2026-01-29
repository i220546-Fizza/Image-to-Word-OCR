# import streamlit as st
# from PIL import Image
# import pytesseract
# from docx import Document
# import io

# # ✅ Correct Tesseract Path
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # -------------------------------
# # Streamlit GUI Settings
# # -------------------------------
# st.set_page_config(
#     page_title="Image to Word OCR Converter",
#     page_icon="📄",
#     layout="centered"
# )

# st.title("📄 Image to Word Converter (OCR)")
# st.write("Upload an image, extract text, and download it as a Word document.")

# st.divider()

# # -------------------------------
# # Upload Image
# # -------------------------------
# uploaded_file = st.file_uploader(
#     "📌 Upload an Image (JPG, PNG)",
#     type=["jpg", "jpeg", "png"]
# )

# # -------------------------------
# # OCR + Word Conversion
# # -------------------------------
# if uploaded_file:
#     image = Image.open(uploaded_file)

#     st.image(image, caption="✅ Uploaded Image Preview", use_column_width=True)

#     if st.button("🔍 Extract Text"):

#         st.info("⏳ Extracting text...")

#         # OCR
#         extracted_text = pytesseract.image_to_string(image)

#         st.success("✅ Text Extracted Successfully!")

#         st.text_area("Extracted Text Output", extracted_text, height=250)

#         # -------------------------------
#         # Create Word Document
#         # -------------------------------
#         doc = Document()
#         doc.add_heading("Extracted Text from Image", level=1)

#         for line in extracted_text.split("\n"):
#             doc.add_paragraph(line)

#         # Save Word file in memory
#         buffer = io.BytesIO()
#         doc.save(buffer)
#         buffer.seek(0)

#         # Download Button
#         st.download_button(
#             label="⬇ Download Word File",
#             data=buffer,
#             file_name="output.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )
import streamlit as st
from PIL import Image
import pytesseract
from docx import Document
import io

# ✅ Correct Tesseract Path
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# -------------------------------
# Streamlit Page Config
# -------------------------------
st.set_page_config(
    page_title="📝 Beige Image to Word OCR",
    page_icon="🖼️",
    layout="wide"
)

# -------------------------------
# Custom CSS for Beige Theme
# -------------------------------
st.markdown("""
<style>
/* Background Beige Gradient */
[data-testid="stAppViewContainer"] {
    background: linear-gradient(135deg, #f5f5dc 0%, #fdf6e3 100%);
    min-height: 100vh;
    padding: 2rem;
    font-family: 'Segoe UI', sans-serif;
}

/* Card container */
.main-card {
    background: rgba(255, 250, 240, 0.95);
    padding: 30px;
    border-radius: 25px;
    box-shadow: 0 15px 35px rgba(0,0,0,0.2);
}

/* Gradient Button in Beige tones */
div.stButton > button:first-child {
    background: linear-gradient(90deg, #d6c4a0, #e6d8b0, #d6c4a0);
    background-size: 200% 200%;
    color: #4b3e2b;
    font-size: 16px;
    font-weight: bold;
    padding: 12px 25px;
    border-radius: 20px;
    border: none;
    cursor: pointer;
    animation: gradientBG 4s ease infinite;
    transition: transform 0.3s;
}
div.stButton > button:first-child:hover {
    transform: scale(1.05);
}
@keyframes gradientBG {
    0%{background-position:0% 50%}
    50%{background-position:100% 50%}
    100%{background-position:0% 50%}
}

/* Text Area */
textarea {
    border-radius: 20px;
    padding: 15px;
    font-size: 16px;
    color: #4b3e2b;
    background-color: #fefaf0;
}

/* File Uploader */
.stFileUploader {
    border: 2px dashed #d6c4a0;
    border-radius: 20px;
    padding: 25px;
    background-color: #fffaf0;
    margin-bottom: 20px;
}

/* Image Styling */
img {
    border-radius: 20px;
    border: 3px solid #e6d8b0;
    box-shadow: 0 10px 25px rgba(0,0,0,0.2);
}

/* Floating Download Button */
.download-btn {
    margin-top: 15px;
    text-align: center;
}
</style>
""", unsafe_allow_html=True)

# -------------------------------
# Title Card
# -------------------------------
st.markdown('<div class="main-card">', unsafe_allow_html=True)
st.title("🖼️ Image to Word OCR")
st.subheader("✨ Upload an image and extract text to Word with a soothing beige theme!")
st.divider()

# -------------------------------
# File Uploader
# -------------------------------
uploaded_file = st.file_uploader(
    "📌 Upload an Image (JPG, PNG)",
    type=["jpg", "jpeg", "png"]
)

# -------------------------------
# OCR + Word Conversion
# -------------------------------
if uploaded_file:
    image = Image.open(uploaded_file)

    # Side-by-side layout
    col1, col2 = st.columns([1,1])

    with col1:
        st.image(image, caption="✅ Uploaded Image Preview", use_column_width=True)

    with col2:
        if st.button("🔍 Extract Text"):
            st.info("⏳ Extracting text...")

            # OCR
            extracted_text = pytesseract.image_to_string(image)
            st.success("✅ Text Extracted Successfully!")

            # Scrollable Text Area
            st.text_area("📝 Extracted Text", extracted_text, height=250)

            # Create Word Document
            doc = Document()
            doc.add_heading("Extracted Text from Image", level=1)
            for line in extracted_text.split("\n"):
                doc.add_paragraph(line)

            # Save Word file in memory
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Floating Download Button
            st.markdown('<div class="download-btn">', unsafe_allow_html=True)
            st.download_button(
                label="⬇ Download Word File",
                data=buffer,
                file_name="extracted_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.markdown('</div>', unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)
